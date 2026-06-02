from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse, JsonResponse
from django.contrib import messages
from django.conf import settings
from django.utils import timezone
from django.core.paginator import Paginator
from django.db import transaction
from django.db.models import Avg, Count, Q
from django.db.utils import DatabaseError, OperationalError, ProgrammingError
from .models import Governorate, Wilayat, Section, Item, Evaluation, Response, ResponseImage, StatisticalRecord
from .activity_options import ACTIVITY_OPTIONS
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Mm, Pt, RGBColor
from io import BytesIO
from pathlib import Path
import json
from datetime import date

EVALUATION_TEAM_OPTIONS = [
    'م. سالم الحارثي (مدير دائرة تطوير نظم سلامة الغذاء)',
    'م. عامر الحبسي (المدير المساعد لدائرة تطوير نظم سلامة الغذاء)',
    'م. ماجد العامري (رئيس قسم نظم سلامة الغذاء)',
    'م. عياض المعولي (رئيس قسم تطبيق نظم سلامة الغذاء)',
    'م. متعب المعمري (رئيس قسم تأهيل المنشآت الغذائية)',
    'د. يوسف أحمد محمد (طبيب بيطري)',
    'م. أحمد المريكي (مفتش صحي)',
]

HACCP_REQUIREMENTS = [
    'سجل استقبال ومخزن المواد الاولية (غذائية)',
    'سجل استقبال ومخزن المواد الاولية (التغليف)',
    'سجل مراقب الجودة (المراقبة اليومية)',
    'سجل مصائد القوارض والحشرات',
    'سجل مراقبة (CCP)',
    'سجل تتبع الإجراءات التصحيحية',
    'سجل تتبع المنتج',
    'سجل التخزين',
    'سجل الارسالية',
    'سجل النظافة',
    'سجل العمال',
    'سجل التدريب',
    'سجل فحص تحليل المختبر',
    'سجلات أخرى',
]


def parse_haccp_documents(value):
    if not value:
        return {}
    try:
        data = json.loads(value)
    except json.JSONDecodeError:
        return {}
    return data if isinstance(data, dict) else {}


def build_haccp_documents_from_request(request):
    documents = {}
    for index, requirement in enumerate(HACCP_REQUIREMENTS):
        documents[str(index)] = {
            'requirement': requirement,
            'exists': request.POST.get(f'haccp_exists_{index}', ''),
            'notes': request.POST.get(f'haccp_notes_{index}', ''),
        }
    return json.dumps(documents, ensure_ascii=False)


def build_haccp_rows(evaluation):
    documents = parse_haccp_documents(evaluation.haccp_documents if evaluation else '')
    rows = []
    for index, requirement in enumerate(HACCP_REQUIREMENTS):
        row = documents.get(str(index), {})
        rows.append({
            'index': index,
            'requirement': row.get('requirement') or requirement,
            'exists': row.get('exists', ''),
            'notes': row.get('notes', ''),
        })
    return rows


def get_wilayats_by_governorate(request):
    """API endpoint to get wilayats for a selected governorate"""
    gov_id = request.GET.get('governorate_id')
    if not gov_id or not str(gov_id).isdigit():
        return JsonResponse({'wilayats': []})
    
    wilayats = Wilayat.objects.filter(governorate_id=int(gov_id)).order_by('name').values('id', 'name')
    return JsonResponse({'wilayats': list(wilayats)})


def get_governorates(request):
    """API endpoint to get all governorates"""
    governorates = Governorate.objects.values('id', 'name')
    return JsonResponse({'governorates': list(governorates)})


def classify_score(score):
    if score >= 86:
        return 'ممتاز، مستوفي للحصول على شهادة ضبط الجودة'
    if score >= 70:
        return 'جيد، مستوفي للحصول على شهادة ضبط الجودة مع وجود فرص للتحسين'
    if score >= 41:
        return 'مقبول، يحتاج تأهيل ومزيد من التحسين'
    return 'ضعيف، إيقاف الإنتاج'


def calculate_score(evaluation):
    # HACCP documents are a checklist only; they are intentionally excluded from scoring.
    totals = evaluation.responses.aggregate(
        total=Count('id', filter=~Q(status='na')),
        compliant=Count('id', filter=Q(status='compliant')),
    )
    total = totals['total'] or 0
    compliant = totals['compliant'] or 0

    score = round((compliant / total) * 100, 2) if total else 0

    evaluation.score = score
    evaluation.classification = classify_score(score)
    evaluation.save()


def is_high_priority(priority):
    return bool(priority and ('عالية' in priority or 'high' in priority.lower()))


def build_evaluation_insights(evaluation, responses, related_statistics):
    responses = list(responses)
    related_statistics = list(related_statistics)
    scored_responses = [r for r in responses if r.status != 'na']
    non_compliant = [r for r in responses if r.status == 'non_compliant']
    compliant_count = len([r for r in scored_responses if r.status == 'compliant'])
    high_priority = [r for r in non_compliant if is_high_priority(r.item.priority)]
    missing_photos = [r for r in non_compliant if not list(r.images.all())]
    missing_corrective = [
        r for r in non_compliant
        if not r.corrective_action.strip() or not r.correction_duration.strip()
    ]

    section_counts = {}
    for response in non_compliant:
        section_title = response.item.section.title
        section_counts[section_title] = section_counts.get(section_title, 0) + 1
    top_sections = [
        {'name': name, 'count': count}
        for name, count in sorted(section_counts.items(), key=lambda item: item[1], reverse=True)[:3]
    ]

    haccp_rows = build_haccp_rows(evaluation)
    missing_haccp = [row for row in haccp_rows if row.get('exists') == 'no']

    if evaluation.score < 41 or len(high_priority) >= 3:
        risk_level = 'حرج'
        risk_class = 'critical'
    elif evaluation.score < 70 or high_priority:
        risk_level = 'مرتفع'
        risk_class = 'high'
    elif evaluation.score < 86 or non_compliant:
        risk_level = 'متوسط'
        risk_class = 'medium'
    else:
        risk_level = 'منخفض'
        risk_class = 'low'

    recommendations = []
    if high_priority:
        recommendations.append('ابدأ بإغلاق البنود عالية الأولوية قبل البنود التشغيلية الأقل أثراً.')
    if missing_corrective:
        recommendations.append('استكمل الإجراء التصحيحي ومدة التصحيح لكل مخالفة قبل اعتماد المتابعة.')
    if missing_photos:
        recommendations.append('أرفق صوراً للمخالفات التي لا تحتوي على توثيق بصري لتسهيل التحقق اللاحق.')
    if missing_haccp:
        recommendations.append('راجع مستندات HACCP وشهادات الجودة الناقصة واربطها بخطة التصحيح.')
    if not recommendations:
        recommendations.append('حافظ على مستوى الالتزام الحالي وجدول متابعة دورية للمنشأة.')

    action_items = sorted(
        non_compliant,
        key=lambda r: (0 if is_high_priority(r.item.priority) else 1, r.item.number),
    )[:5]

    return {
        'risk_level': risk_level,
        'risk_class': risk_class,
        'total_items': len(scored_responses),
        'compliant_count': compliant_count,
        'non_compliant_count': len(non_compliant),
        'high_priority_count': len(high_priority),
        'missing_photos_count': len(missing_photos),
        'missing_corrective_count': len(missing_corrective),
        'missing_haccp_count': len(missing_haccp),
        'top_sections': top_sections,
        'recommendations': recommendations,
        'action_items': action_items,
        'related_records_count': len(related_statistics),
        'latest_related_record': related_statistics[0] if related_statistics else None,
    }


def parse_iso_date(value):
    if not value:
        return None
    try:
        return date.fromisoformat(value)
    except ValueError:
        return None


def parse_float(value):
    if value in (None, ''):
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def parse_coordinate(value):
    parsed = parse_float(value)
    return round(parsed, 7) if parsed is not None else None


def apply_evaluation_filters(queryset, filters):
    query = filters.get('q', '').strip()
    license_number = filters.get('license_number', '').strip()
    cr_number = filters.get('cr_number', '').strip()
    governorate_id = filters.get('governorate', '').strip()
    classification = filters.get('classification', '').strip()
    status = filters.get('status', '').strip()
    date_from = parse_iso_date(filters.get('date_from', '').strip())
    date_to = parse_iso_date(filters.get('date_to', '').strip())
    min_score = parse_float(filters.get('min_score', '').strip())
    max_score = parse_float(filters.get('max_score', '').strip())

    if query:
        queryset = queryset.filter(
            Q(facility_name__icontains=query)
            | Q(contact_name__icontains=query)
            | Q(license_number__icontains=query)
            | Q(cr_number__icontains=query)
        )

    if license_number:
        queryset = queryset.filter(license_number__icontains=license_number)

    if cr_number:
        queryset = queryset.filter(cr_number__icontains=cr_number)

    if governorate_id.isdigit():
        queryset = queryset.filter(governorate_id=int(governorate_id))

    if classification:
        queryset = queryset.filter(classification__icontains=classification)

    if status == 'draft':
        queryset = queryset.filter(is_draft=True)
    elif status == 'final':
        queryset = queryset.filter(is_draft=False)

    if date_from:
        queryset = queryset.filter(visit_date__gte=date_from)

    if date_to:
        queryset = queryset.filter(visit_date__lte=date_to)

    if min_score is not None:
        queryset = queryset.filter(score__gte=min_score)

    if max_score is not None:
        queryset = queryset.filter(score__lte=max_score)

    return queryset


def get_evaluation_summary(queryset):
    data = queryset.aggregate(
        total=Count('id'),
        avg_score=Avg('score'),
        excellent=Count('id', filter=Q(score__gte=86)),
        needs_attention=Count('id', filter=Q(score__lt=70)),
    )
    return {
        'total': data['total'] or 0,
        'avg_score': round(data['avg_score'] or 0, 1),
        'excellent': data['excellent'] or 0,
        'needs_attention': data['needs_attention'] or 0,
    }


def as_percent(part, total):
    return round((part / total) * 100, 1) if total else 0


def build_statistics_context():
    evaluations = Evaluation.objects.all()
    records = StatisticalRecord.objects.select_related('report').all()
    evaluation_totals = evaluations.aggregate(
        total_reports=Count('id'),
        avg_score=Avg('score'),
        excellent=Count('id', filter=Q(score__gte=86)),
        good=Count('id', filter=Q(score__gte=70, score__lt=86)),
        acceptable=Count('id', filter=Q(score__gte=41, score__lt=70)),
        weak=Count('id', filter=Q(score__lt=41)),
    )
    record_totals = records.aggregate(
        total_records=Count('id'),
        linked_records=Count('id', filter=Q(report__isnull=False)),
    )
    total_reports = evaluation_totals['total_reports'] or 0
    total_records = record_totals['total_records'] or 0
    linked_records = record_totals['linked_records'] or 0

    score_bands = [
        {'label': 'ممتاز', 'count': evaluation_totals['excellent'] or 0, 'class': 'excellent'},
        {'label': 'جيد', 'count': evaluation_totals['good'] or 0, 'class': 'good'},
        {'label': 'مقبول', 'count': evaluation_totals['acceptable'] or 0, 'class': 'acceptable'},
        {'label': 'ضعيف', 'count': evaluation_totals['weak'] or 0, 'class': 'weak'},
    ]
    for band in score_bands:
        band['percent'] = as_percent(band['count'], total_reports)

    by_governorate = list(
        evaluations
        .values('governorate__name')
        .annotate(count=Count('id'), avg_score=Avg('score'))
        .order_by('-count', 'governorate__name')[:10]
    )
    for row in by_governorate:
        row['name'] = row.pop('governorate__name') or 'غير محدد'
        row['avg_score'] = round(row['avg_score'] or 0, 1)
        row['percent'] = as_percent(row['count'], total_reports)

    by_activity = list(
        evaluations
        .exclude(activity_type='')
        .values('activity_type')
        .annotate(count=Count('id'), avg_score=Avg('score'))
        .order_by('-count', 'activity_type')[:10]
    )
    for row in by_activity:
        row['avg_score'] = round(row['avg_score'] or 0, 1)
        row['percent'] = as_percent(row['count'], total_reports)

    record_categories = list(
        records
        .exclude(category='')
        .values('category')
        .annotate(count=Count('id'))
        .order_by('-count', 'category')[:10]
    )
    for row in record_categories:
        row['percent'] = as_percent(row['count'], total_records)

    recent_records = records.order_by('-visit_date', '-updated_at')[:20]
    common_non_compliant_sections = list(
        Response.objects
        .filter(status='non_compliant')
        .values('item__section__title')
        .annotate(count=Count('id'))
        .order_by('-count', 'item__section__title')[:5]
    )
    total_common_non_compliant = sum(item['count'] for item in common_non_compliant_sections)
    for row in common_non_compliant_sections:
        row['name'] = row.pop('item__section__title') or 'غير محدد'
        row['percent'] = as_percent(row['count'], total_common_non_compliant)

    risk_reports = Evaluation.objects.select_related('governorate').filter(score__lt=70).order_by('score', '-visit_date')[:5]
    weak_count = evaluation_totals['weak'] or 0
    needs_attention_count = weak_count + (evaluation_totals['acceptable'] or 0)
    smart_alerts = []
    if total_reports and needs_attention_count:
        smart_alerts.append(f'{as_percent(needs_attention_count, total_reports)}% من التقارير تحتاج متابعة تحسين أو تدخل.')
    if common_non_compliant_sections:
        smart_alerts.append(f'أكثر محور تتكرر فيه المخالفات: {common_non_compliant_sections[0]["name"]}.')
    if total_records and total_records - linked_records:
        smart_alerts.append('توجد سجلات إحصائية غير مرتبطة بتقارير، وهذا يقلل دقة القراءة التاريخية.')
    if not smart_alerts:
        smart_alerts.append('مؤشرات النظام مستقرة ولا توجد تنبيهات رئيسية حالياً.')

    return {
        'summary': {
            'total_reports': total_reports,
            'avg_score': round(evaluation_totals['avg_score'] or 0, 1),
            'total_records': total_records,
            'linked_records': linked_records,
            'unlinked_records': total_records - linked_records,
        },
        'score_bands': score_bands,
        'by_governorate': by_governorate,
        'by_activity': by_activity,
        'record_categories': record_categories,
        'recent_records': recent_records,
        'common_non_compliant_sections': common_non_compliant_sections,
        'risk_reports': risk_reports,
        'smart_alerts': smart_alerts,
    }


def resolve_location(gov_id, wilayat_id):
    governorate = Governorate.objects.filter(id=gov_id).first() if gov_id else None
    wilayat = Wilayat.objects.filter(id=wilayat_id).first() if wilayat_id else None

    if wilayat and governorate and wilayat.governorate_id != governorate.id:
        wilayat = None
    if wilayat and not governorate:
        governorate = wilayat.governorate

    return governorate, wilayat


def evaluation_list(request):
    filters = {
        'q': request.GET.get('q', '').strip(),
        'license_number': request.GET.get('license_number', '').strip(),
        'cr_number': request.GET.get('cr_number', '').strip(),
        'governorate': request.GET.get('governorate', '').strip(),
        'classification': request.GET.get('classification', '').strip(),
        'status': request.GET.get('status', '').strip(),
        'date_from': request.GET.get('date_from', '').strip(),
        'date_to': request.GET.get('date_to', '').strip(),
        'min_score': request.GET.get('min_score', '').strip(),
        'max_score': request.GET.get('max_score', '').strip(),
    }

    error_message = ''
    query_params = request.GET.copy()
    query_params.pop('page', None)
    try:
        evaluations = apply_evaluation_filters(
            Evaluation.objects.select_related('governorate', 'wilayat'),
            filters,
        )
        results_count = evaluations.count()
        summary = get_evaluation_summary(evaluations)
        paginator = Paginator(evaluations, 25)
        page_obj = paginator.get_page(request.GET.get('page'))
        governorates = Governorate.objects.all()
    except (OperationalError, ProgrammingError, DatabaseError):
        evaluations = Evaluation.objects.none()
        page_obj = None
        results_count = 0
        summary = {'total': 0, 'avg_score': 0, 'excellent': 0, 'needs_attention': 0}
        governorates = Governorate.objects.none()
        error_message = 'تعذر استجلاب التقارير من قاعدة البيانات. تأكد من اتصال DATABASE_URL وتطبيق migrations المطلوبة.'

    return render(request, 'inspections/evaluation_list.html', {
        'evaluations': page_obj.object_list if page_obj else evaluations,
        'page_obj': page_obj,
        'filters': filters,
        'results_count': results_count,
        'querystring': query_params.urlencode(),
        'summary': summary,
        'governorates': governorates,
        'db_error_message': error_message,
    })


def statistics_dashboard(request):
    error_message = ''
    try:
        context = build_statistics_context()
    except (OperationalError, ProgrammingError, DatabaseError):
        context = {
            'summary': {
                'total_reports': 0,
                'avg_score': 0,
                'total_records': 0,
                'linked_records': 0,
                'unlinked_records': 0,
            },
            'score_bands': [],
            'by_governorate': [],
            'by_activity': [],
            'record_categories': [],
            'recent_records': [],
            'common_non_compliant_sections': [],
            'risk_reports': [],
            'smart_alerts': [],
        }
        error_message = 'تعذر تحميل الإحصائيات من قاعدة البيانات. تأكد من تطبيق migrations واستيراد ملف Excel عند الحاجة.'

    context['db_error_message'] = error_message
    return render(request, 'inspections/statistics_dashboard.html', context)


def get_reports(request):
    filters = {
        'q': request.GET.get('q', '').strip(),
        'license_number': request.GET.get('license_number', '').strip(),
        'cr_number': request.GET.get('cr_number', '').strip(),
        'governorate': request.GET.get('governorate', '').strip(),
        'classification': request.GET.get('classification', '').strip(),
        'status': request.GET.get('status', '').strip(),
        'date_from': request.GET.get('date_from', '').strip(),
        'date_to': request.GET.get('date_to', '').strip(),
        'min_score': request.GET.get('min_score', '').strip(),
        'max_score': request.GET.get('max_score', '').strip(),
    }

    try:
        evaluations = apply_evaluation_filters(Evaluation.objects.all(), filters)[:500]
    except (OperationalError, ProgrammingError, DatabaseError):
        return JsonResponse({
            'reports': [],
            'count': 0,
            'error': 'Database schema/connection is not ready. Run migrations for the connected database.',
        }, status=503)

    data = [
        {
            'id': evaluation.pk,
            'facility_name': evaluation.facility_name,
            'license_number': evaluation.license_number,
            'cr_number': evaluation.cr_number,
            'visit_date': evaluation.visit_date,
            'score': evaluation.score,
            'classification': evaluation.classification,
            'is_draft': evaluation.is_draft,
        }
        for evaluation in evaluations
    ]
    return JsonResponse({'reports': data, 'count': len(data)})


@transaction.atomic
def save_evaluation_from_request(request, evaluation=None):
    sections = Section.objects.prefetch_related('items').all()
    extra_fields = [
        'shift_count',
        'workers_per_shift',
        'total_factory_workers',
        'supervisors_per_shift',
        'total_supervisors',
        'product_types',
        'distribution_scope',
        'actual_daily_production_rate',
        'permitted_daily_production_rate',
        'water_source',
        'final_product_storage_area',
        'location_url',
        'haccp_manual',
        'iso_22000_certificate',
        'haccp_certificate',
        'other_quality_certificate',
    ]
    
    # Resolve location safely even if posted IDs are invalid.
    gov_id = request.POST.get('governorate')
    wilayat_id = request.POST.get('wilayat')
    governorate, wilayat = resolve_location(gov_id, wilayat_id)

    if evaluation is None:
        evaluation = Evaluation.objects.create(
            facility_name=request.POST.get('facility_name', ''),
            activity_type=request.POST.get('activity_type', ''),
            license_number=request.POST.get('license_number', ''),
            cr_number=request.POST.get('cr_number', ''),
            governorate=governorate,
            wilayat=wilayat,
            location_latitude=parse_coordinate(request.POST.get('location_latitude')),
            location_longitude=parse_coordinate(request.POST.get('location_longitude')),
            contact_name=request.POST.get('contact_name', ''),
            contact_phone=request.POST.get('contact_phone', ''),
            visit_date=request.POST.get('visit_date'),
            evaluation_team=request.POST.get('evaluation_team', ''),
            haccp_documents=build_haccp_documents_from_request(request),
            is_draft=False,
            **{field: request.POST.get(field, '') for field in extra_fields},
        )
    else:
        evaluation.facility_name = request.POST.get('facility_name', '')
        evaluation.activity_type = request.POST.get('activity_type', '')
        evaluation.license_number = request.POST.get('license_number', '')
        evaluation.cr_number = request.POST.get('cr_number', '')
        evaluation.governorate = governorate
        evaluation.wilayat = wilayat
        evaluation.location_latitude = parse_coordinate(request.POST.get('location_latitude'))
        evaluation.location_longitude = parse_coordinate(request.POST.get('location_longitude'))
        evaluation.contact_name = request.POST.get('contact_name', '')
        evaluation.contact_phone = request.POST.get('contact_phone', '')
        evaluation.visit_date = request.POST.get('visit_date')
        evaluation.evaluation_team = request.POST.get('evaluation_team', '')
        evaluation.haccp_documents = build_haccp_documents_from_request(request)
        evaluation.is_draft = False
        for field in extra_fields:
            setattr(evaluation, field, request.POST.get(field, ''))
        evaluation.save()

    for section in sections:
        for item in section.items.all():
            status = request.POST.get(f'status_{item.id}', '').strip()

            if not status:
                response = Response.objects.filter(evaluation=evaluation, item=item).first()
                if response and response.images.exists():
                    continue
                if response:
                    response.delete()
                continue

            response, created = Response.objects.update_or_create(
                evaluation=evaluation,
                item=item,
                defaults={
                    'status': status,
                    'notes': request.POST.get(f'notes_{item.id}', ''),
                    'corrective_action': request.POST.get(f'corrective_{item.id}', ''),
                    'correction_duration': request.POST.get(f'duration_{item.id}', ''),
                }
            )

            for image in request.FILES.getlist(f'images_{item.id}'):
                ResponseImage.objects.create(response=response, image=image)

    calculate_score(evaluation)
    return evaluation


def evaluation_form(request):
    db_error_message = ''
    try:
        # Evaluate querysets eagerly so DB schema errors are caught here.
        sections = list(Section.objects.prefetch_related('items').all())
        governorates = list(Governorate.objects.all())
    except (OperationalError, ProgrammingError, DatabaseError):
        sections = []
        governorates = []
        db_error_message = 'تعذر تحميل أقسام التقييم من قاعدة البيانات. تأكد من تطبيق migrations وتهيئة البيانات.'
    
    if request.method == 'POST':
        evaluation = save_evaluation_from_request(request)
        messages.success(request, 'تم حفظ التقييم وحساب النتيجة بنجاح.')
        return redirect('report_detail', evaluation.pk)

    return render(request, 'inspections/evaluation_form.html', {
        'sections': sections,
        'evaluation': None,
        'responses_map': {},
        'governorates': governorates,
        'wilayats': Wilayat.objects.none(),
        'activity_options': ACTIVITY_OPTIONS,
        'evaluation_team_options': EVALUATION_TEAM_OPTIONS,
        'haccp_rows': build_haccp_rows(None),
        'today': timezone.localdate(),
        'db_error_message': db_error_message,
    })


def evaluation_edit(request, pk):
    evaluation = get_object_or_404(Evaluation, pk=pk)
    db_error_message = ''
    try:
        # Evaluate querysets eagerly so DB schema errors are caught here.
        sections = list(Section.objects.prefetch_related('items').all())
        governorates = list(Governorate.objects.all())
        wilayats = list(evaluation.governorate.wilayats.all()) if evaluation.governorate else []
        responses_map = {
            r.item_id: r
            for r in evaluation.responses.select_related('item').prefetch_related('images')
        }
    except (OperationalError, ProgrammingError, DatabaseError):
        sections = []
        governorates = []
        wilayats = []
        responses_map = {}
        db_error_message = 'تعذر تحميل بيانات التقييم من قاعدة البيانات. تأكد من تطبيق migrations وتهيئة البيانات.'

    if request.method == 'POST':
        evaluation = save_evaluation_from_request(request, evaluation)
        messages.success(request, 'تم تعديل التقييم بنجاح.')
        return redirect('report_detail', evaluation.pk)

    return render(request, 'inspections/evaluation_form.html', {
        'sections': sections,
        'evaluation': evaluation,
        'responses_map': responses_map,
        'governorates': governorates,
        'wilayats': wilayats,
        'activity_options': ACTIVITY_OPTIONS,
        'evaluation_team_options': EVALUATION_TEAM_OPTIONS,
        'haccp_rows': build_haccp_rows(evaluation),
        'db_error_message': db_error_message,
    })


def report_detail(request, pk):
    evaluation = get_object_or_404(
        Evaluation.objects.select_related('governorate', 'wilayat'),
        pk=pk,
    )
    responses = evaluation.responses.select_related('item', 'item__section').prefetch_related('images').all()
    related_statistics = StatisticalRecord.objects.filter(
        Q(report=evaluation) | Q(facility_name__iexact=evaluation.facility_name)
    ).order_by('-visit_date', '-updated_at')
    smart_insights = build_evaluation_insights(evaluation, responses, related_statistics)
    return render(request, 'inspections/report_detail.html', {
        'evaluation': evaluation,
        'responses': responses,
        'haccp_rows': build_haccp_rows(evaluation),
        'related_statistics': related_statistics,
        'smart_insights': smart_insights,
    })


def evaluation_delete(request, pk):
    evaluation = get_object_or_404(Evaluation, pk=pk)
    if request.method == 'POST':
        evaluation.delete()
        messages.success(request, 'تم حذف التقرير بنجاح.')
        return redirect('evaluation_list')
    return render(request, 'inspections/confirm_delete.html', {'evaluation': evaluation})


def shade_cell(cell, fill):
    cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill)))


def set_table_rtl(table):
    table._tbl.tblPr.append(parse_xml(r'<w:bidiVisual {} />'.format(nsdecls('w'))))


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(str(text or ''))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_table_header(row, fill='0F6B4B'):
    for cell in row.cells:
        shade_cell(cell, fill)
        set_cell_text(cell, cell.text, bold=True, color='FFFFFF')


def add_section_heading(doc, title, fill='0F6B4B'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_text(cell, title, bold=True, color='FFFFFF')
    doc.add_paragraph()


def add_evaluation_team_block(doc, evaluation_team):
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    set_table_rtl(table)

    header_cell = table.rows[0].cells[0]
    shade_cell(header_cell, '155E75')
    set_cell_text(header_cell, 'فريق التقييم', bold=True, color='FFFFFF')

    body_cell = table.rows[1].cells[0]
    shade_cell(body_cell, 'EEF7FA')
    body_cell.text = ''
    team_members = [member.strip() for member in str(evaluation_team or '').splitlines() if member.strip()]
    if not team_members:
        team_members = ['']

    for index, member in enumerate(team_members):
        paragraph = body_cell.paragraphs[0] if index == 0 else body_cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(member)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string('155E75')

    doc.add_paragraph()


def get_report_header_path():
    base_dir = Path(settings.BASE_DIR)
    candidates = [
        base_dir / 'static' / 'images' / 'report_header.png',
        base_dir / 'inspections' / 'static' / 'inspections' / 'report_header.png',
        base_dir.parent.parent / 'final' / 'static' / 'images' / 'report_header.png',
    ]
    for path in candidates:
        if path.exists():
            return path
    return None


def add_report_header(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    section.top_margin = Mm(8)
    section.bottom_margin = Mm(10)

    header_path = get_report_header_path()
    if not header_path:
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run().add_picture(str(header_path), width=section.page_width)


def get_report_signature_path():
    base_dir = Path(settings.BASE_DIR)
    candidates = [
        base_dir / 'static' / 'images' / 'report_signature.png',
        base_dir / 'inspections' / 'static' / 'inspections' / 'report_signature.png',
        base_dir.parent.parent / 'final' / 'static' / 'images' / 'report_signature.png',
    ]
    for path in candidates:
        if path.exists():
            return path
    return None


def add_report_signature(doc):
    signature_path = get_report_signature_path()
    if not signature_path:
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(signature_path), width=Inches(4.6))


def export_word(request, pk):
    evaluation = get_object_or_404(
        Evaluation.objects.select_related('governorate', 'wilayat'),
        pk=pk,
    )
    non_compliant_responses = (
        evaluation.responses
        .filter(status='non_compliant')
        .select_related('item', 'item__section')
        .prefetch_related('images')
        .order_by('item__section__order', 'item__number')
    )
    doc = Document()
    add_report_header(doc)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    title = doc.add_heading('تقرير تقييم منشأة غذائية', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor.from_string('0F6B4B')

    score_fill = '138A4C' if evaluation.score >= 86 else '2F855A' if evaluation.score >= 70 else 'D9A441' if evaluation.score >= 41 else 'B42318'
    score_table = doc.add_table(rows=1, cols=2)
    score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    score_table.style = 'Table Grid'
    score_cells = score_table.rows[0].cells
    shade_cell(score_cells[0], score_fill)
    shade_cell(score_cells[1], score_fill)
    set_cell_text(score_cells[0], f'{evaluation.score}%', bold=True, color='FFFFFF')
    set_cell_text(score_cells[1], evaluation.classification, bold=True, color='FFFFFF')

    add_section_heading(doc, 'بيانات المنشأة')
    info = [
        ('اسم المنشأة', evaluation.facility_name),
        ('نوع النشاط', evaluation.activity_type),
        ('رقم الرخصة', evaluation.license_number),
        ('رقم السجل التجاري', evaluation.cr_number),
        ('المحافظة', evaluation.governorate.name if evaluation.governorate else ''),
        ('الولاية', evaluation.wilayat.name if evaluation.wilayat else ''),
        ('الموقع المباشر', evaluation.location_url),
        ('تاريخ الزيارة', str(evaluation.visit_date)),
        ('عدد الورديات', evaluation.shift_count),
        ('عدد العاملين في الوردية', evaluation.workers_per_shift),
        ('مجموع عدد العاملين المصنع', evaluation.total_factory_workers),
        ('عدد المشرفين في الوردية', evaluation.supervisors_per_shift),
        ('مجموع عدد المشرفين', evaluation.total_supervisors),
        ('أنواع المنتجات', evaluation.product_types),
        ('التوزيع (محلي/تصدير)', evaluation.distribution_scope),
        ('معدل الإنتاج اليومي الفعلي', evaluation.actual_daily_production_rate),
        ('معدل الإنتاج اليومي المسموح', evaluation.permitted_daily_production_rate),
        ('مصدر الماء المستخدم', evaluation.water_source),
        ('إجمالي مساحة مخزن المنتج النهائي', evaluation.final_product_storage_area),
        ('دليل الهاسب للشركة إن وجد', evaluation.haccp_manual),
        ('آيزو 22000', evaluation.iso_22000_certificate),
        ('الهاسب', evaluation.haccp_certificate),
        ('أخرى', evaluation.other_quality_certificate),
        ('النتيجة', f'{evaluation.score}%'),
        ('الوضع العام للمنشأة', evaluation.classification),
    ]

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for label, value in info:
        row = table.add_row().cells
        shade_cell(row[1], 'E6F4EE')
        set_cell_text(row[0], value)
        set_cell_text(row[1], label, bold=True)

    add_section_heading(doc, 'البنود غير المستوفية والإجراءات التصحيحية', fill='B42318')

    if not non_compliant_responses:
        ok = doc.add_paragraph('لا توجد بنود غير مستوفية.')
        ok.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        current_section = None
        for r in non_compliant_responses:
            if current_section != r.item.section_id:
                current_section = r.item.section_id
            add_section_heading(doc, f'\u200e{r.item.section.order}. {r.item.section.title}', fill='155E75')

            item_table = doc.add_table(rows=1, cols=6)
            item_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            item_table.style = 'Table Grid'
            set_table_rtl(item_table)
            hdr = item_table.rows[0].cells
            headers = ['البند', 'الأولوية', 'الملاحظات', 'الإجراء التصحيحي', 'مدة التصحيح', 'الصور']
            for cell, header in zip(hdr, headers):
                set_cell_text(cell, header, bold=True)
            style_table_header(item_table.rows[0], fill='0F6B4B')

            row = item_table.add_row().cells
            set_cell_text(row[0], f'\u200e{r.item.number} - {r.item.text}')
            set_cell_text(row[1], r.item.priority)
            set_cell_text(row[2], r.notes)
            set_cell_text(row[3], r.corrective_action)
            set_cell_text(row[4], r.correction_duration)

            image_paragraph = row[5].paragraphs[0]
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            images_added = 0
            for image in r.images.all():
                try:
                    import urllib.request, io
                    img_url = image.image.url
                    if img_url.startswith('http'):
                        with urllib.request.urlopen(img_url) as resp:
                            img_data = io.BytesIO(resp.read())
                        image_paragraph.add_run().add_picture(img_data, width=Inches(1.1))
                    else:
                        image_paragraph.add_run().add_picture(image.image.path, width=Inches(1.1))
                    images_added += 1
                except Exception:
                    continue
            if not images_added:
                set_cell_text(row[5], 'لا توجد صور')

            doc.add_paragraph()

    add_section_heading(doc, 'مستندات الهاسب وشهادات أنظمة الجودة')
    table3 = doc.add_table(rows=1, cols=4)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    table3.style = 'Table Grid'
    hdr = table3.rows[0].cells
    hdr[0].text = 'الملاحظات'
    hdr[1].text = 'لا'
    hdr[2].text = 'نعم'
    hdr[3].text = 'المتطلبات'
    style_table_header(table3.rows[0])

    for haccp_row in build_haccp_rows(evaluation):
        row = table3.add_row().cells
        set_cell_text(row[0], haccp_row['notes'])
        set_cell_text(row[1], 'نعم' if haccp_row['exists'] == 'no' else '')
        set_cell_text(row[2], 'نعم' if haccp_row['exists'] == 'yes' else '')
        set_cell_text(row[3], haccp_row['requirement'])

    add_evaluation_team_block(doc, evaluation.evaluation_team)
    add_report_signature(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=evaluation_report_{evaluation.pk}.docx'
    return response
